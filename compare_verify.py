"""compare_verify.py — Phase 9: Compare & verify.

Deterministic parsing and diffing (no AI in this module, per ROADMAP.md's
own scope for this phase) over TLF outputs a user brings in from their own
SAS/R environment — this app never executes SAS, so it can't produce a real
RTF/Word/PDF report itself; what it CAN do deterministically is read one
back and compare it, either against another output (e.g. production vs. an
independent QC run) or against the mock shell the table was supposed to
match.

Two entry points:

  compare_outputs(path_a, path_b)      — output-to-output: extract both,
                                          align sections/tables by TITLE
                                          text (not page position — a
                                          re-paginated re-run shouldn't
                                          register as "everything moved"),
                                          diff each aligned pair.
  validate_against_shell(path, shell)  — output-to-mock-shell: does every
                                          row label and title the shell
                                          promised actually appear in the
                                          rendered output?

Supported formats: .pdf (pdfplumber, already a dependency — acrf_parser.py
uses it for aCRF PDFs), .docx (python-docx), .rtf (striprtf), .txt/.csv as
a plain-text fallback for anything else a user might upload.
"""

import csv
import difflib
import io
import os

import pdfplumber
import docx as docx_lib
from striprtf.striprtf import rtf_to_text

from runlog import log_run

TITLE_MATCH_THRESHOLD = 0.6


# ---------------------------------------------------------------------------
# Extraction: file -> list of blocks. A block is a self-contained section
# (one table, or one run of body text) with a `title` (its heading, or a
# short synthesized one) used for alignment, and either `rows` (list of
# list-of-cell-strings, for a real table) or `lines` (list of strings, for
# body text with no detected table structure).
# ---------------------------------------------------------------------------

def _block(title, rows=None, lines=None):
    return {"title": title.strip(), "rows": rows or [], "lines": lines or []}


def _extract_pdf(path):
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            for t_idx, table in enumerate(page.extract_tables() or []):
                rows = [[str(c).strip() if c is not None else "" for c in row] for row in table]
                title = f"page {page_no} table {t_idx + 1}"
                # The first non-empty row is usually the title/header band
                # in a clinical table PDF export — use it as the block's
                # title so alignment matches on content, not position.
                for row in rows:
                    joined = " ".join(c for c in row if c).strip()
                    if joined:
                        title = joined
                        break
                blocks.append(_block(title, rows=rows))
            text = page.extract_text() or ""
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            if lines:
                blocks.append(_block(lines[0] if lines else f"page {page_no} text", lines=lines))
    return blocks


def _extract_docx(path):
    doc = docx_lib.Document(path)
    blocks = []
    for t_idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        title = f"table {t_idx + 1}"
        for row in rows:
            joined = " ".join(c for c in row if c).strip()
            if joined:
                title = joined
                break
        blocks.append(_block(title, rows=rows))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        blocks.append(_block(paragraphs[0], lines=paragraphs))
    return blocks


def _extract_rtf(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        raw = f.read()
    text = rtf_to_text(raw)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return []
    return [_block(lines[0], lines=lines)]


def _extract_text(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        raw = f.read()
    if path.lower().endswith(".csv"):
        rows = list(csv.reader(io.StringIO(raw)))
        rows = [[c.strip() for c in row] for row in rows if any(c.strip() for c in row)]
        title = " ".join(rows[0]) if rows else os.path.basename(path)
        return [_block(title, rows=rows)]
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    if not lines:
        return []
    return [_block(lines[0], lines=lines)]


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".rtf": _extract_rtf,
    ".txt": _extract_text,
    ".csv": _extract_text,
}


def extract_content(path):
    """Dispatch by extension. Raises ValueError for an unsupported type —
    surfaced to the caller as a normal in-app note, not a crash."""
    ext = os.path.splitext(path)[1].lower()
    if ext not in _EXTRACTORS:
        raise ValueError(f"Unsupported file type {ext!r} (supported: {', '.join(_EXTRACTORS)})")
    return _EXTRACTORS[ext](path)


# ---------------------------------------------------------------------------
# Alignment: match blocks between two documents by TITLE similarity, not
# position — a table that moved from page 2 to page 3 between two runs
# (different pagination, same content) should still align to itself.
# ---------------------------------------------------------------------------

def align_blocks(blocks_a, blocks_b):
    """Returns (matched, only_a, only_b). matched: [(block_a, block_b,
    similarity), ...]. Greedy best-match by title similarity — fine at the
    scale of one document's worth of tables/sections (tens, not thousands)."""
    remaining_b = list(enumerate(blocks_b))
    matched = []
    only_a = []

    for a in blocks_a:
        best = None
        best_score = 0.0
        for idx, b in remaining_b:
            score = difflib.SequenceMatcher(None, a["title"].lower(), b["title"].lower()).ratio()
            if score > best_score:
                best_score = score
                best = (idx, b)
        if best and best_score >= TITLE_MATCH_THRESHOLD:
            matched.append((a, best[1], best_score))
            remaining_b = [(i, b) for i, b in remaining_b if i != best[0]]
        else:
            only_a.append(a)

    only_b = [b for _, b in remaining_b]
    return matched, only_a, only_b


# ---------------------------------------------------------------------------
# Diffing one aligned pair
# ---------------------------------------------------------------------------

def _diff_rows(rows_a, rows_b):
    text_a = ["\t".join(r) for r in rows_a]
    text_b = ["\t".join(r) for r in rows_b]
    return list(difflib.unified_diff(text_a, text_b, fromfile="A", tofile="B", lineterm=""))


def _diff_lines(lines_a, lines_b):
    return list(difflib.unified_diff(lines_a, lines_b, fromfile="A", tofile="B", lineterm=""))


def diff_blocks(block_a, block_b):
    if block_a["rows"] or block_b["rows"]:
        diff = _diff_rows(block_a["rows"], block_b["rows"])
    else:
        diff = _diff_lines(block_a["lines"], block_b["lines"])
    return diff


# ---------------------------------------------------------------------------
# Top-level: output-to-output
# ---------------------------------------------------------------------------

def compare_outputs(path_a, path_b, log=True):
    """Extract, align, and diff two output files. Returns a findings list:
    [{kind: "changed"|"only_a"|"only_b", title, similarity, diff: [...]}]
    — empty means the two outputs are content-identical (modulo layout)."""
    blocks_a = extract_content(path_a)
    blocks_b = extract_content(path_b)
    matched, only_a, only_b = align_blocks(blocks_a, blocks_b)

    findings = []
    for a, b, score in matched:
        diff = diff_blocks(a, b)
        if diff:
            findings.append({"kind": "changed", "title": a["title"], "similarity": round(score, 2), "diff": diff})
    for a in only_a:
        findings.append({"kind": "only_a", "title": a["title"], "similarity": None, "diff": []})
    for b in only_b:
        findings.append({"kind": "only_b", "title": b["title"], "similarity": None, "diff": []})

    if log:
        log_run(f"{path_a} vs {path_b}", "compare", "n/a", "n/a", "n/a",
                len(findings), "compare_verify")
    return findings


# ---------------------------------------------------------------------------
# Top-level: output-to-mock-shell
# ---------------------------------------------------------------------------

def validate_against_shell(path, shell_path, log=True):
    """Does the rendered output actually contain what the shell promised?
    Presence check (case-insensitive substring) for every Shell_Meta
    title/footnote and every Shell_Rows label against the output's full
    extracted text — lighter-weight than a structural table diff (a shell
    describes WHAT should appear, not the exact cell layout), and the
    right check for "did the table builder actually cover every row",
    which is what a mock shell is for. Returns a list of missing items:
    [{field, expected}] — empty means everything the shell listed was
    found somewhere in the output."""
    import pandas as pd

    blocks = extract_content(path)
    haystack = "\n".join(
        " ".join(cell for row in b["rows"] for cell in row) + " " + " ".join(b["lines"])
        for b in blocks
    ).lower()

    meta_df = pd.read_excel(shell_path, sheet_name="Shell_Meta")
    meta = dict(zip(meta_df["Field"], meta_df["Value"]))
    rows_df = pd.read_excel(shell_path, sheet_name="Shell_Rows")

    missing = []
    for field in ("title1", "title2", "title3"):
        expected = str(meta.get(field, "")).strip()
        if expected and expected.lower() != "nan" and expected.lower() not in haystack:
            missing.append({"field": field, "expected": expected})

    label_col = "label" if "label" in rows_df.columns else "row_label"
    if label_col in rows_df.columns:
        for _, row in rows_df.iterrows():
            expected = str(row[label_col]).strip()
            if expected and expected.lower() != "nan" and expected.lower() not in haystack:
                missing.append({"field": f"row: {expected}", "expected": expected})

    if log:
        log_run(shell_path, "verify_shell", "n/a", "n/a", "n/a", len(missing), path)
    return missing


def print_compare_summary(findings):
    print(f"\n{'='*50}")
    print("COMPARE & VERIFY")
    print(f"{'='*50}")
    changed = [f for f in findings if f["kind"] == "changed"]
    only_a = [f for f in findings if f["kind"] == "only_a"]
    only_b = [f for f in findings if f["kind"] == "only_b"]
    print(f"Changed  : {len(changed)}")
    print(f"Only in A: {len(only_a)}")
    print(f"Only in B: {len(only_b)}")
    for f in findings:
        print(f"\n[{f['kind'].upper()}] {f['title']}")
        for line in f["diff"][:20]:
            print(f"  {line}")

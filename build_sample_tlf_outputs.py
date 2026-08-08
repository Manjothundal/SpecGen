"""build_sample_tlf_outputs.py — generate sample rendered TLF outputs
(PDF/DOCX/RTF) for compare_verify.py to exercise and test against.

This app never executes SAS/R, so it can't produce a REAL rendered table —
these are small synthetic stand-ins for what a user's own SAS/R run would
produce, matching sample_shell_demographics.xlsx's table 14.1.1 layout
closely enough to test output-to-mock-shell validation, and built as an
A/B pair (one changed cell) to test output-to-output comparison. Mirrors
the existing build_sample_acrf.py / build_sample_protocol.py pattern:
generate once, commit the output as a fixture.
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

import docx as docx_lib

STYLES = getSampleStyleSheet()


def build_pdf(path, age_mean="52.3"):
    doc = SimpleDocTemplate(path, pagesize=letter)
    title = Paragraph(
        "Table 14.1.1 Summary of Demographic and Baseline Characteristics",
        STYLES["Title"],
    )
    data = [
        ["Table 14.1.1 Summary of Demographic and Baseline Characteristics", "", ""],
        ["Parameter", "Placebo (N=50)", "Drug A 50mg (N=50)"],
        ["Age (years)", f"Mean: {age_mean}", "Mean: 54.0"],
        ["Sex, n (%)", "Male: 25 (50%)", "Male: 27 (54%)"],
    ]
    table = Table(data)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
    doc.build([title, table])


def build_docx(path):
    d = docx_lib.Document()
    d.add_heading("Table 14.1.1 Summary of Demographic and Baseline Characteristics", level=1)
    table = d.add_table(rows=3, cols=3)
    values = [
        ["Parameter", "Placebo (N=50)", "Drug A 50mg (N=50)"],
        ["Age (years)", "Mean: 52.3", "Mean: 54.0"],
        ["Sex, n (%)", "Male: 25 (50%)", "Male: 27 (54%)"],
    ]
    for r, row_vals in enumerate(values):
        for c, val in enumerate(row_vals):
            table.rows[r].cells[c].text = val
    d.save(path)


def build_rtf(path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(r"""{\rtf1\ansi\deff0
{\fonttbl{\f0 Courier;}}
\f0\fs20
Table 14.1.1 Summary of Demographic and Baseline Characteristics\par
Safety Population\par
\par
Parameter          Placebo (N=50)     Drug A 50mg (N=50)\par
Age (years)         Mean: 52.3          Mean: 54.0\par
Sex, n (%)          Male: 25 (50%)      Male: 27 (54%)\par
}
""")


if __name__ == "__main__":
    build_pdf("sample_output_a.pdf", age_mean="52.3")
    build_pdf("sample_output_b.pdf", age_mean="53.8")
    build_docx("sample_output.docx")
    build_rtf("sample_output.rtf")
    print("Wrote sample_output_a.pdf, sample_output_b.pdf, sample_output.docx, sample_output.rtf")
